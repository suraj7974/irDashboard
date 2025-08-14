import sys
import re
import argparse
import os
import json
import subprocess
from pathlib import Path
from datetime import datetime
import logging
from typing import Dict, List, Optional

# PDF and Document processing libraries
try:
    import PyPDF2
    import pdfplumber
    import fitz  # PyMuPDF
    import pandas as pd
    import tabula
    from docx import Document
    
    # macOS alternatives for .doc files
    try:
        import pypandoc
        PANDOC_AVAILABLE = True
    except ImportError:
        PANDOC_AVAILABLE = False
    
    try:
        import docx2txt
        DOCX2TXT_AVAILABLE = True
    except ImportError:
        DOCX2TXT_AVAILABLE = False
    
    PDF_LIBRARIES_AVAILABLE = True
except ImportError as e:
    PDF_LIBRARIES_AVAILABLE = False
    MISSING_LIBRARIES = str(e)

# Global constants
CURRENT_USER = "vanshrajchauhan"
CURRENT_DATE = "2025-08-08"
CURRENT_TIME = "17:28:25"

# Setup logging
def setup_logging(log_dir: str):
    """Setup logging for bulk processing."""
    log_file = Path(log_dir) / f"conversion_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8'),
            logging.StreamHandler()
        ]
    )
    return logging.getLogger(__name__)

def check_dependencies():
    """Check if required libraries are available."""
    if not PDF_LIBRARIES_AVAILABLE:
        print("Error: Required libraries not found.")
        print("Please install the required libraries:")
        print("pip install PyPDF2 pdfplumber PyMuPDF pandas tabula-py python-docx")
        print("For .doc support on macOS:")
        print("brew install antiword pandoc")
        print("pip install pypandoc docx2txt")
        print("Also install Java for tabula-py: brew install openjdk")
        print(f"Missing: {MISSING_LIBRARIES}")
        return False
    return True

def get_unicode(text):
    """Convert input text to Unicode string."""
    if isinstance(text, str):
        return text
    try:
        return text.decode('utf-8')
    except (UnicodeDecodeError, AttributeError):
        try:
            return text.decode('unicode_escape')
        except (UnicodeDecodeError, AttributeError):
            try:
                return text.decode('latin-1')
            except (UnicodeDecodeError, AttributeError):
                return str(text)

def krutidev_to_unicode(kru_text):
    """
    Convert KrutiDev text to Unicode Devanagari.
    
    Args:
        kru_text (str): Text in KrutiDev encoding
        
    Returns:
        str: Text converted to Unicode Devanagari
    """
    
    # KrutiDev to Unicode mapping table
    k2u_mappings = [
        # Special characters and numbers
        ('\u00f1', '\u0970'),  # ñ -> ॰
        ('Q+Z', 'QZ+'),
        ('sas', 'sa'),
        ('aa', 'a'),
        (')Z', '\u0930\u094d\u0926\u094d\u0927'),  # )Z -> र्द्ध
        ('ZZ', 'Z'),
        ('\u2018', '"'),  # ' -> "
        ('\u2019', '"'),  # ' -> "
        ('\u201c', "'"),  # " -> '
        ('\u201d', "'"),  # " -> '
        
        # Numbers
        ('\u00e5', '\u0966'),  # å -> ०
        ('\u0192', '\u0967'),  # ƒ -> १
        ('\u201e', '\u0968'),  # „ -> २
        ('\u2026', '\u0969'),  # … -> ③
        ('\u2020', '\u096a'),  # † -> ④
        ('\u2021', '\u096b'),  # ‡ -> ⑤
        ('\u02c6', '\u096c'),  # ˆ -> ⑥
        ('\u2030', '\u096d'),  # ‰ -> ⑦
        ('\u0160', '\u096e'),  # Š -> ⑧
        ('\u2039', '\u096f'),  # ‹ -> ⑨
        
        # Extended characters
        ('\u00b6+', '\u095e\u094d'),  # ¶+ -> फ़्
        ('d+', '\u0958'),  # d+ -> क़
        ('[+k', '\u0959'),  # [+k -> ख़
        ('[+', '\u0959\u094d'),  # [+ -> ख़्
        ('x+', '\u095a'),  # x+ -> ग़
        ('T+', '\u091c\u093c\u094d'),  # T+ -> ज़्
        ('t+', '\u095b'),  # t+ -> ज़
        ('M+', '\u095c'),  # M+ -> ड़
        ('<+', '\u095d'),  # <+ -> ढ़
        ('Q+', '\u095e'),  # Q+ -> फ़
        (';+', '\u095f'),  # ;+ -> य़
        ('j+', '\u0931'),  # j+ -> ऱ
        ('u+', '\u0929'),  # u+ -> ऩ
        
        # Complex consonants
        ('\u00d9k', '\u0924\u094d\u0924'),  # Ùk -> त्त
        ('\u00d9', '\u0924\u094d\u0924\u094d'),  # Ù -> त्त्
        ('\u00e4', '\u0915\u094d\u0924'),  # ä -> क्त
        ('\u2013', '\u0926\u0943'),  # – -> दृ
        ('\u2014', '\u0915\u0943'),  # — -> कृ
        ('\u00e9', '\u0928\u094d\u0928'),  # é -> न्न
        ('\u2122', '\u0928\u094d\u0928\u094d'),  # ™ -> न्न्
        ('=kk', '=k'),
        ('f=k', 'f='),
        ('\u00e0', '\u0939\u094d\u0928'),  # à -> ह्न
        ('\u00e1', '\u0939\u094d\u092f'),  # á -> ह्य
        ('\u00e2', '\u0939\u0943'),  # â -> हृ
        ('\u00e3', '\u0939\u094d\u092e'),  # ã -> ह्म
        ('\u00baz', '\u0939\u094d\u0930'),  # ºz -> ह्र
        ('\u00ba', '\u0939\u094d'),  # º -> ह्
        ('\u00ed', '\u0926\u094d\u0926'),  # í -> द्द
        ('{k', '\u0915\u094d\u0937'),  # {k -> क्ष
        ('{', '\u0915\u094d\u0937\u094d'),  # { -> क्ष्
        ('=', '\u0924\u094d\u0930'),  # = -> त्र
        ('\u00ab', '\u0924\u094d\u0930\u094d'),  # « -> त्र्
        
        # Conjuncts with य
        ('N\u00ee', '\u091b\u094d\u092f'),  # Nî -> छ्य
        ('V\u00ee', '\u091f\u094d\u092f'),  # Vî -> ट्य
        ('B\u00ee', '\u0920\u094d\u092f'),  # Bî -> ठ्य
        ('M\u00ee', '\u0921\u094d\u092f'),  # Mî -> ड्य
        ('<\u00ee', '\u0922\u094d\u092f'),  # <î -> ढ्य
        ('|', '\u0926\u094d\u092f'),  # | -> द्य
        
        # Special conjuncts
        ('K', '\u091c\u094d\u091e'),  # K -> ज्ञ
        ('}', '\u0926\u094d\u0935'),  # } -> द्व
        ('J', '\u0936\u094d\u0930'),  # J -> श्र
        
        # र् conjuncts
        ('V\u00aa', '\u091f\u094d\u0930'),  # Vª -> ट्र
        ('M\u00aa', '\u0921\u094d\u0930'),  # Mª -> ड्र
        ('<\u00aa\u00aa', '\u0922\u094d\u0930'),  # <ªª -> ढ्र
        ('N\u00aa', '\u091b\u094d\u0930'),  # Nª -> छ्र
        ('\u00d8', '\u0915\u094d\u0930'),  # Ø -> क्र
        ('\u00dd', '\u092b\u094d\u0930'),  # Ý -> फ्र
        ('nzZ', '\u0930\u094d\u0926\u094d\u0930'),  # nzZ -> र्द्र
        ('\u00e6', '\u0926\u094d\u0930'),  # æ -> द्र
        ('\u00e7', '\u092a\u094d\u0930'),  # ç -> प्र
        ('\u00c1', '\u092a\u094d\u0930'),  # Á -> प्र
        ('xz', '\u0917\u094d\u0930'),  # xz -> ग्र
        
        # र with vowels
        ('#', '\u0930\u0941'),  # # -> रु
        (':', '\u0930\u0942'),  # : -> रू
        
        # Independent vowels
        ('v\u201a', '\u0911'),  # v‚ -> ऑ
        ('vks', '\u0913'),  # vks -> ओ
        ('vkS', '\u0914'),  # vkS -> औ
        ('vk', '\u0906'),  # vk -> आ
        ('v', '\u0905'),  # v -> अ
        ('b\u00b1', '\u0908\u0902'),  # b± -> ईं
        ('\u00c3', '\u0908'),  # Ã -> ई
        ('bZ', '\u0908'),  # bZ -> ई
        ('b', '\u0907'),  # b -> इ
        ('m', '\u0909'),  # m -> उ
        ('\u00c5', '\u090a'),  # Å -> ऊ
        (',s', '\u0910'),  # ,s -> ऐ
        (',', '\u090f'),  # , -> ए
        ('_', '\u090b'),  # _ -> ऋ
        
        # Consonants
        ('\u00f4', '\u0915\u094d\u0915'),  # ô -> क्क
        ('d', '\u0915'),  # d -> क
        ('Dk', '\u0915'),  # Dk -> क
        ('D', '\u0915\u094d'),  # D -> क्
        ('[k', '\u0916'),  # [k -> ख
        ('[', '\u0916\u094d'),  # [ -> ख्
        ('x', '\u0917'),  # x -> ग
        ('Xk', '\u0917'),  # Xk -> ग
        ('X', '\u0917\u094d'),  # X -> ग्
        ('\u00c4', '\u0918'),  # Ä -> घ
        ('?k', '\u0918'),  # ?k -> घ
        ('?', '\u0918\u094d'),  # ? -> घ्
        ('\u00b3', '\u0919'),  # ³ -> ङ
        ('pkS', '\u091a\u0948'),  # pkS -> चै
        ('p', '\u091a'),  # p -> च
        ('Pk', '\u091a'),  # Pk -> च
        ('P', '\u091a\u094d'),  # P -> च्
        ('N', '\u091b'),  # N -> छ
        ('t', '\u091c'),  # t -> ज
        ('Tk', '\u091c'),  # Tk -> ज
        ('T', '\u091c\u094d'),  # T -> ज्
        ('>', '\u091d'),  # > -> झ
        ('\u00f7', '\u091d\u094d'),  # ÷ -> झ्
        ('\u00a5', '\u091e'),  # ¥ -> ञ
        ('\u00ea', '\u091f\u094d\u091f'),  # ê -> ट्ट
        ('\u00eb', '\u091f\u094d\u0920'),  # ë -> ट्ठ
        ('V', '\u091f'),  # V -> ट
        ('B', '\u0920'),  # B -> ठ
        ('\u00ec', '\u0921\u094d\u0921'),  # ì -> ड्ड
        ('\u00ef', '\u0921\u094d\u0922'),  # ï -> ड्ढ
        ('M', '\u0921'),  # M -> ड
        ('<', '\u0922'),  # < -> ढ
        ('.k', '\u0923'),  # .k -> ण
        ('.', '\u0923\u094d'),  # . -> ण्
        ('r', '\u0924'),  # r -> त
        ('Rk', '\u0924'),  # Rk -> त
        ('R', '\u0924\u094d'),  # R -> त्
        ('Fk', '\u0925'),  # Fk -> थ
        ('F', '\u0925\u094d'),  # F -> थ्
        (')', '\u0926\u094d\u0927'),  # ) -> द्ध
        ('n', '\u0926'),  # n -> द
        ('/k', '\u0927'),  # /k -> ध
        ('/', '\u0927\u094d'),  # / -> ध्
        ('\u00cb', '\u0927\u094d'),  # Ë -> ध्
        ('\u00e8', '\u0927'),  # è -> ध
        ('u', '\u0928'),  # u -> न
        ('Uk', '\u0928'),  # Uk -> न
        ('U', '\u0928\u094d'),  # U -> न्
        ('i', '\u092a'),  # i -> प
        ('Ik', '\u092a'),  # Ik -> प
        ('I', '\u092a\u094d'),  # I -> प्
        ('Q', '\u092b'),  # Q -> फ
        ('\u00b6', '\u092b\u094d'),  # ¶ -> फ्
        ('c', '\u092c'),  # c -> ब
        ('Ck', '\u092c'),  # Ck -> ब
        ('C', '\u092c\u094d'),  # C -> ब्
        ('Hk', '\u092d'),  # Hk -> भ
        ('H', '\u092d\u094d'),  # H -> भ्
        ('e', '\u092e'),  # e -> म
        ('Ek', '\u092e'),  # Ek -> म
        ('E', '\u092e\u094d'),  # E -> म्
        (';', '\u092f'),  # ; -> य
        ('\u00b8', '\u092f\u094d'),  # ¸ -> य्
        ('j', '\u0930'),  # j -> र
        ('y', '\u0932'),  # y -> ल
        ('Yk', '\u0932'),  # Yk -> ल
        ('Y', '\u0932\u094d'),  # Y -> ल्
        ('G', '\u0933'),  # G -> ळ
        ('o', '\u0935'),  # o -> व
        ('Ok', '\u0935'),  # Ok -> व
        ('O', '\u0935\u094d'),  # O -> व्
        ("'k", '\u0936'),  # 'k -> श
        ("'", '\u0936\u094d'),  # ' -> श्
        ('"k', '\u0937'),  # "k -> ष
        ('"', '\u0937\u094d'),  # " -> ष्
        ('l', '\u0938'),  # l -> स
        ('Lk', '\u0938'),  # Lk -> स
        ('L', '\u0938\u094d'),  # L -> स्
        ('g', '\u0939'),  # g -> ह
        
        # Vowel signs and modifiers
        ('\u00c8', '\u0940\u0902'),  # È -> ीं
        ('saz', '\u094d\u0930\u0947\u0902'),  # saz -> ्रें
        ('z', '\u094d\u0930'),  # z -> ्र
        ('\u00dc', '\u0936\u094d'),  # Ü -> श्
        ('\u201a', '\u0949'),  # ‚ -> ॉ
        ('kas', '\u094b\u0902'),  # kas -> ों
        ('ks', '\u094b'),  # ks -> ो
        ('kS', '\u094c'),  # kS -> ौ
        ('\u00a1k', '\u093e\u0901'),  # ¡k -> ाँ
        ('ak', 'k\u0902'),  # ak -> kं
        ('k', '\u093e'),  # k -> ा
        ('ah', '\u0940\u0902'),  # ah -> ीं
        ('h', '\u0940'),  # h -> ी
        ('aq', '\u0941\u0902'),  # aq -> ुं
        ('q', '\u0941'),  # q -> ु
        ('aw', '\u0942\u0902'),  # aw -> ूं
        ('\u00a1w', '\u0942\u0901'),  # ¡w -> ूँ
        ('w', '\u0942'),  # w -> ू
        ('`', '\u0943'),  # ` -> ृ
        ('\u0300', '\u0943'),  # ̀ -> ृ
        ('as', '\u0947\u0902'),  # as -> ें
        ('\u00b1s', 's\u00b1'),  # ±s -> s±
        ('s', '\u0947'),  # s -> े
        ('aS', '\u0948\u0902'),  # aS -> ैं
        ('S', '\u0948'),  # S -> ै
        ('a\u00aa', '\u094d\u0930\u0902'),  # aª -> ्रं
        ('\u00aa', '\u094d\u0930'),  # ª -> ्र
        ('fa', '\u0902f'),  # fa -> ंf
        ('a', '\u0902'),  # a -> ं
        ('\u00a1', '\u0901'),  # ¡ -> ँ
        
        # Punctuation and symbols
        ('%', ':'),  # % -> :
        ('W', '\u0945'),  # W -> ॅ
        ('\u2022', '\u093d'),  # • -> ऽ
        ('\u00b7', '\u093d'),  # · -> ऽ
        ('\u2219', '\u093d'),  # ∙ -> ऽ
        ('~j', '\u094d\u0930'),  # ~j -> ्र
        ('~', '\u094d'),  # ~ -> ्
        ('\\', '?'),  # \ -> ?
        ('+', '\u093c'),  # + -> ़
        ('^', '\u2018'),  # ^ -> '
        ('*', '\u2019'),  # * -> '
        ('\u00de', '\u201c'),  # Þ -> "
        ('\u00df', '\u201d'),  # ß -> "
        ('(', ';'),  # ( -> ;
        ('\u00bc', '('),  # ¼ -> (
        ('\u00bd', ')'),  # ½ -> )
        ('\u00bf', '{'),  # ¿ -> {
        ('\u00c0', '}'),  # À -> }
        ('\u00be', '='),  # ¾ -> =
        ('A', '\u0964'),  # A -> ।
        ('-', '.'),  # - -> .
        ('&', '-'),  # & -> -
        ('&', '\u00b5'),  # & -> µ
        ('\u03bc', '-'),  # μ -> -
        ('\u0152', '\u0970'),  # Œ -> ॰
        (']', ','),  # ] -> ,
        ('~ ', '\u094d '),  # ~ -> ् 
        ('@', '/'),  # @ -> /
        ('\u00ae', '\u0948\u0902'),  # ® -> ैं
    ]
    
    # Unicode vowel signs for processing
    unicode_unattached_vowel_signs = [
        '\u093e', '\u093f', '\u0940', '\u0941', '\u0942', '\u0943',
        '\u0947', '\u0948', '\u094b', '\u094c', '\u0902', '\u0903',
        '\u0901', '\u0945'
    ]
    
    # Convert to Unicode if needed
    kru_text = get_unicode(kru_text)
    
    if not kru_text:
        return ""
    
    # Pre-processing fixes
    kru_text = kru_text.replace(' \u00aa', '\u00aa')
    kru_text = kru_text.replace(' ~j', '~j')
    kru_text = kru_text.replace(' z', 'z')
    
    # Handle misplaced en-dash and em-dash
    krutidev_consonants = ['d', '[k', 'x', '?k', '\u00b3', 'p', 'N', 't', '>', '\u00a5', 'V', 'B', 'M', '<', '.k', 'r', 'Fk', 'n', '/k', 'u', 'i', 'Q', 'c', 'Hk', 'e', ';', 'j', 'y', 'G', 'o', "'k", '"k', 'l', 'g']
    krutidev_unattached_vowel_signs = ['k', 'f', 'h', 'q', 'w', '`', 's', 'S', 'ks', 'kS', 'a', '%', '\u00a1', 'W']
    
    misplaced = re.compile('[\u2014\u2013]')
    for m in misplaced.finditer(kru_text):
        index = m.start()
        if index < len(kru_text) - 1 and kru_text[m.start() + 1] not in krutidev_consonants + krutidev_unattached_vowel_signs:
            kru_text = kru_text[:index] + '&' + kru_text[index + 1:]
    
    # Apply all mappings
    for old, new in k2u_mappings:
        kru_text = kru_text.replace(old, new)
    
    # Special post-processing fixes
    kru_text = kru_text.replace('\u00b1', 'Z\u0902')  # ± -> Zं
    kru_text = kru_text.replace('\u00c6', '\u0930\u094df')  # Æ -> र्f
    
    # Fix f + ? -> ? + ि
    misplaced = re.search('f(.?)', kru_text)
    while misplaced:
        char = misplaced.group(1)
        kru_text = kru_text.replace('f' + char, char + '\u093f')
        misplaced = re.search('f(.?)', kru_text)
    
    kru_text = kru_text.replace('\u00c7', 'fa')  # Ç -> fa
    kru_text = kru_text.replace('\u00af', 'fa')  # ¯ -> fa
    kru_text = kru_text.replace('\u00c9', '\u0930\u094dfa')  # É -> र्fa
    
    # Fix fa? -> ? + िं
    misplaced = re.search('fa(.?)', kru_text)
    while misplaced:
        char = misplaced.group(1)
        kru_text = kru_text.replace('fa' + char, char + '\u093f\u0902')
        misplaced = re.search('fa(.?)', kru_text)
    
    kru_text = kru_text.replace('\u00ca', '\u0940Z')  # Ê -> ीZ
    
    # Fix ि् + ? -> ् + ? + ि
    misplaced = re.search('\u093f\u094d(.?)', kru_text)
    while misplaced:
        char = misplaced.group(1)
        kru_text = kru_text.replace('\u093f\u094d' + char, '\u094d' + char + '\u093f')
        misplaced = re.search('\u093f\u094d(.?)', kru_text)
    
    kru_text = kru_text.replace('\u094dZ', 'Z')  # ् + Z -> Z
    
    # Fix र + ् placement
    misplaced = re.search('(.?)Z', kru_text)
    while misplaced:
        char = misplaced.group(1)
        if char + 'Z' in kru_text:
            index_r_halant = kru_text.index(char + 'Z')
            while index_r_halant >= 0 and kru_text[index_r_halant] in ['\u0905', '\u0906', '\u0907', '\u0908', '\u0909', '\u090a', '\u090f', '\u0910', '\u0913', '\u0914', '\u093e', '\u093f', '\u0940', '\u0941', '\u0942', '\u0943', '\u0947', '\u0948', '\u094b', '\u094c', '\u0902', '\u0903', '\u0901', '\u0945']:
                index_r_halant -= 1
                if index_r_halant >= 0:
                    char = kru_text[index_r_halant] + char
                else:
                    break
            kru_text = kru_text.replace(char + 'Z', '\u0930\u094d' + char)
        misplaced = re.search('(.?)Z', kru_text)
    
    # Clean up illegal characters before matras
    for matra in unicode_unattached_vowel_signs:
        kru_text = kru_text.replace(' ' + matra, matra)
        kru_text = kru_text.replace(',' + matra, matra + ',')
        kru_text = kru_text.replace('\u094d' + matra, matra)
    
    # Final cleanup
    kru_text = kru_text.replace('\u094d\u094d\u0930', '\u094d\u0930')  # ् + ् + र -> ् + र
    kru_text = kru_text.replace('\u094d\u0930\u094d', '\u0930\u094d')  # ् + र + ् -> र + ्
    kru_text = kru_text.replace('\u094d\u094d', '\u094d')  # ् + ् -> ्
    kru_text = kru_text.replace('\u094d ', ' ')  # Remove trailing halant
    
    return kru_text

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF using multiple methods."""
    try:
        # Try pdfplumber first
        with pdfplumber.open(pdf_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text
    except Exception:
        pass
    
    try:
        # Try PyMuPDF
        doc = fitz.open(pdf_path)
        text = ""
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text() + "\n"
        doc.close()
        if text.strip():
            return text
    except Exception:
        pass
    
    try:
        # Try PyPDF2
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        return text
    except Exception:
        pass
    
    return ""

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        doc = Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text
    except Exception as e:
        return ""

def extract_text_from_doc_macos(doc_path: str) -> str:
    """Extract text from DOC file on macOS using multiple methods."""
    
    # Method 1: Try antiword
    try:
        result = subprocess.run(['antiword', doc_path], 
                              capture_output=True, text=True, encoding='utf-8')
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except Exception:
        pass
    
    # Method 2: Try pandoc
    if PANDOC_AVAILABLE:
        try:
            text = pypandoc.convert_file(doc_path, 'plain')
            if text and text.strip():
                return text
        except Exception:
            pass
    
    # Method 3: Try docx2txt (sometimes works for .doc files too)
    if DOCX2TXT_AVAILABLE:
        try:
            import docx2txt
            text = docx2txt.process(doc_path)
            if text and text.strip():
                return text
        except Exception:
            pass
    
    # Method 4: Try converting with LibreOffice if available
    try:
        result = subprocess.run([
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            '--headless', '--convert-to', 'txt', '--outdir', '/tmp', doc_path
        ], capture_output=True, text=True)
        
        if result.returncode == 0:
            # Read the converted text file
            doc_name = Path(doc_path).stem
            txt_file = f"/tmp/{doc_name}.txt"
            if os.path.exists(txt_file):
                with open(txt_file, 'r', encoding='utf-8') as f:
                    text = f.read()
                os.remove(txt_file)  # Clean up
                return text
    except Exception:
        pass
    
    return ""

def extract_tables_from_pdf(pdf_path: str) -> List[Dict]:
    """Extract tables from PDF with page-wise organization."""
    tables = []
    
    # Try pdfplumber
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_tables = page.extract_tables()
                for table_num, table in enumerate(page_tables, 1):
                    if table:
                        headers = table[0] if table else []
                        rows = table[1:] if len(table) > 1 else []
                        
                        table_data = {
                            'page': page_num,
                            'table_number': table_num,
                            'headers': [krutidev_to_unicode(str(cell) if cell else '') for cell in headers],
                            'rows': []
                        }
                        
                        for row in rows:
                            converted_row = [krutidev_to_unicode(str(cell) if cell else '') for cell in row]
                            table_data['rows'].append(converted_row)
                        
                        tables.append(table_data)
    except Exception:
        pass
    
    # Try tabula if no tables found
    if not tables:
        try:
            dfs = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True)
            for table_num, df in enumerate(dfs, 1):
                if not df.empty:
                    table_data = {
                        'table_number': table_num,
                        'headers': [krutidev_to_unicode(str(col)) for col in df.columns],
                        'rows': []
                    }
                    
                    for _, row in df.iterrows():
                        converted_row = [krutidev_to_unicode(str(cell) if pd.notna(cell) else '') for cell in row]
                        table_data['rows'].append(converted_row)
                    
                    tables.append(table_data)
        except Exception:
            pass
    
    return tables

def extract_tables_from_docx(docx_path: str) -> List[Dict]:
    """Extract tables from DOCX file."""
    tables = []
    try:
        doc = Document(docx_path)
        for table_num, table in enumerate(doc.tables, 1):
            if table.rows:
                # Get headers from first row
                headers = [krutidev_to_unicode(cell.text) for cell in table.rows[0].cells]
                
                table_data = {
                    'table_number': table_num,
                    'headers': headers,
                    'rows': []
                }
                
                # Get data rows
                for row in table.rows[1:]:
                    converted_row = [krutidev_to_unicode(cell.text) for cell in row.cells]
                    table_data['rows'].append(converted_row)
                
                tables.append(table_data)
    except Exception:
        pass
    
    return tables

def extract_pagewise_data_from_pdf(pdf_path: str, extract_tables: bool = True) -> List[Dict]:
    """Extract page-wise data from PDF including text and tables."""
    pages_data = []
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text from current page
                page_text = page.extract_text()
                if page_text:
                    converted_text = krutidev_to_unicode(page_text)
                else:
                    converted_text = ""
                
                # Extract tables from current page
                page_tables = []
                if extract_tables:
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        if table:
                            headers = table[0] if table else []
                            rows = table[1:] if len(table) > 1 else []
                            
                            table_data = {
                                'table_number': table_num,
                                'headers': [krutidev_to_unicode(str(cell) if cell else '') for cell in headers],
                                'rows': []
                            }
                            
                            for row in rows:
                                converted_row = [krutidev_to_unicode(str(cell) if cell else '') for cell in row]
                                table_data['rows'].append(converted_row)
                            
                            page_tables.append(table_data)
                
                # Create page data structure
                page_data = {
                    'page_number': page_num,
                    'text_content': converted_text,
                    'text_length': len(converted_text),
                    'tables_count': len(page_tables),
                    'tables': page_tables
                }
                
                pages_data.append(page_data)
    
    except Exception as e:
        # Fallback to non-page-wise extraction
        text_content = extract_text_from_pdf(pdf_path)
        converted_text = krutidev_to_unicode(text_content) if text_content else ""
        tables = extract_tables_from_pdf(pdf_path) if extract_tables else []
        
        pages_data = [{
            'page_number': 1,
            'text_content': converted_text,
            'text_length': len(converted_text),
            'tables_count': len(tables),
            'tables': tables
        }]
    
    return pages_data

def find_files_recursively(root_path: str, extensions: List[str]) -> List[Dict]:
    """Find all files with specified extensions recursively."""
    files_found = []
    root_path = Path(root_path)
    
    for file_path in root_path.rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in extensions:
            # Extract path components for organization
            relative_path = file_path.relative_to(root_path)
            path_parts = list(relative_path.parts)
            
            # Assuming structure: area_committee/year/file
            area_committee = path_parts[0] if len(path_parts) > 0 else "unknown"
            year = path_parts[1] if len(path_parts) > 1 else "unknown"
            
            file_info = {
                'file_path': str(file_path),
                'relative_path': str(relative_path),
                'filename': file_path.name,
                'extension': file_path.suffix.lower(),
                'area_committee': area_committee,
                'year': year,
                'path_parts': path_parts[:-1],  # All parts except filename
                'file_size': file_path.stat().st_size if file_path.exists() else 0
            }
            files_found.append(file_info)
    
    return files_found

def process_single_file(file_info: Dict, extract_tables: bool = True, pagewise: bool = False) -> Optional[Dict]:
    """Process a single file and extract its content."""
    file_path = file_info['file_path']
    extension = file_info['extension']
    
    try:
        # Extract text and tables based on file type
        if extension == '.pdf':
            if pagewise:
                pages_data = extract_pagewise_data_from_pdf(file_path, extract_tables)
                total_text_length = sum(page['text_length'] for page in pages_data)
                total_tables = sum(page['tables_count'] for page in pages_data)
                
                file_data = {
                    'file_metadata': {
                        'original_filename': file_info['filename'],
                        'file_extension': extension,
                        'area_committee': file_info['area_committee'],
                        'year': file_info['year'],
                        'relative_path': file_info['relative_path'],
                        'full_path': file_path,
                        'file_size_bytes': file_info['file_size'],
                        'processing_date': f"{CURRENT_DATE}T{CURRENT_TIME}Z",
                        'processed_by': CURRENT_USER,
                        'total_pages': len(pages_data)
                    },
                    'pages': pages_data,
                    'summary': {
                        'total_text_length': total_text_length,
                        'total_tables': total_tables,
                        'pages_with_text': sum(1 for page in pages_data if page['text_length'] > 0),
                        'pages_with_tables': sum(1 for page in pages_data if page['tables_count'] > 0),
                    }
                }
            else:
                text_content = extract_text_from_pdf(file_path)
                tables = extract_tables_from_pdf(file_path) if extract_tables else []
                converted_text = krutidev_to_unicode(text_content) if text_content else ""
                
                file_data = {
                    'file_metadata': {
                        'original_filename': file_info['filename'],
                        'file_extension': extension,
                        'area_committee': file_info['area_committee'],
                        'year': file_info['year'],
                        'relative_path': file_info['relative_path'],
                        'full_path': file_path,
                        'file_size_bytes': file_info['file_size'],
                        'processing_date': f"{CURRENT_DATE}T{CURRENT_TIME}Z",
                        'processed_by': CURRENT_USER
                    },
                    'content': {
                        'text_content': converted_text,
                        'text_length': len(converted_text),
                        'tables_count': len(tables),
                        'tables': tables
                    }
                }
        
        elif extension == '.docx':
            text_content = extract_text_from_docx(file_path)
            tables = extract_tables_from_docx(file_path) if extract_tables else []
            converted_text = krutidev_to_unicode(text_content) if text_content else ""
            
            file_data = {
                'file_metadata': {
                    'original_filename': file_info['filename'],
                    'file_extension': extension,
                    'area_committee': file_info['area_committee'],
                    'year': file_info['year'],
                    'relative_path': file_info['relative_path'],
                    'full_path': file_path,
                    'file_size_bytes': file_info['file_size'],
                    'processing_date': f"{CURRENT_DATE}T{CURRENT_TIME}Z",
                    'processed_by': CURRENT_USER
                },
                'content': {
                    'text_content': converted_text,
                    'text_length': len(converted_text),
                    'tables_count': len(tables),
                    'tables': tables
                }
            }
        
        elif extension == '.doc':
            text_content = extract_text_from_doc_macos(file_path)
            converted_text = krutidev_to_unicode(text_content) if text_content else ""
            
            file_data = {
                'file_metadata': {
                    'original_filename': file_info['filename'],
                    'file_extension': extension,
                    'area_committee': file_info['area_committee'],
                    'year': file_info['year'],
                    'relative_path': file_info['relative_path'],
                    'full_path': file_path,
                    'file_size_bytes': file_info['file_size'],
                    'processing_date': f"{CURRENT_DATE}T{CURRENT_TIME}Z",
                    'processed_by': CURRENT_USER
                },
                'content': {
                    'text_content': converted_text,
                    'text_length': len(converted_text),
                    'tables_count': 0,
                    'tables': []
                }
            }
        
        else:
            return None
        
        # Add extraction summary
        if 'content' in file_data:
            file_data['extraction_summary'] = {
                'has_text': bool(file_data['content']['text_content'].strip()),
                'has_tables': file_data['content']['tables_count'] > 0,
                'extraction_method': {
                    'text': 'multi-method' if extension == '.pdf' else 'direct',
                    'tables': 'multi-method' if extension == '.pdf' else 'direct'
                }
            }
        
        return file_data
        
    except Exception as e:
        logging.error(f"Error processing file {file_path}: {str(e)}")
        return None

def bulk_convert_files(input_dir: str, output_dir: str, extract_tables: bool = True, 
                      supported_extensions: List[str] = None, pagewise: bool = False) -> Dict:
    """Convert all supported files in input directory to JSON."""
    
    if supported_extensions is None:
        supported_extensions = ['.pdf', '.docx', '.doc']
    
    input_path = Path(input_dir)
    output_path = Path(output_dir)
    
    # Create output directory if it doesn't exist
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Setup logging
    logger = setup_logging(str(output_path))
    
    # Find all files
    logger.info(f"Scanning for files in: {input_dir}")
    files_found = find_files_recursively(input_dir, supported_extensions)
    logger.info(f"Found {len(files_found)} files to process")
    
    # Process files
    processed_files = []
    failed_files = []
    
    for i, file_info in enumerate(files_found, 1):
        logger.info(f"Processing {i}/{len(files_found)}: {file_info['relative_path']}")
        
        try:
            # Process the file
            file_data = process_single_file(file_info, extract_tables, pagewise)
            
            if file_data:
                # Create output filename
                area_committee = file_info['area_committee']
                year = file_info['year']
                filename_without_ext = Path(file_info['filename']).stem
                
                # Create area committee directory in output
                area_output_dir = output_path / area_committee / year
                area_output_dir.mkdir(parents=True, exist_ok=True)
                
                # Save JSON file
                suffix = "_pagewise" if pagewise else "_converted"
                json_filename = f"{filename_without_ext}{suffix}.json"
                json_path = area_output_dir / json_filename
                
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(file_data, f, ensure_ascii=False, indent=2)
                
                # Calculate stats
                if 'content' in file_data:
                    text_length = file_data['content']['text_length']
                    tables_count = file_data['content']['tables_count']
                elif 'summary' in file_data:
                    text_length = file_data['summary']['total_text_length']
                    tables_count = file_data['summary']['total_tables']
                else:
                    text_length = 0
                    tables_count = 0
                
                processed_files.append({
                    'source_file': file_info['relative_path'],
                    'output_file': str(json_path.relative_to(output_path)),
                    'text_length': text_length,
                    'tables_count': tables_count
                })
                
                logger.info(f"Successfully processed: {file_info['relative_path']}")
                
            else:
                failed_files.append(file_info['relative_path'])
                logger.warning(f"Failed to process: {file_info['relative_path']}")
                
        except Exception as e:
            failed_files.append(file_info['relative_path'])
            logger.error(f"Error processing {file_info['relative_path']}: {str(e)}")
    
    # Create summary report
    summary = {
        'conversion_summary': {
            'input_directory': input_dir,
            'output_directory': output_dir,
            'processing_date': f"{CURRENT_DATE}T{CURRENT_TIME}Z",
            'processed_by': CURRENT_USER,
            'total_files_found': len(files_found),
            'successfully_processed': len(processed_files),
            'failed_files': len(failed_files),
            'supported_extensions': supported_extensions,
            'tables_extracted': extract_tables,
            'pagewise_extraction': pagewise,
            'macos_compatibility': True
        },
        'processed_files': processed_files,
        'failed_files': failed_files,
        'file_distribution': {},
        'statistics': {
            'total_text_length': sum(f['text_length'] for f in processed_files),
            'total_tables': sum(f['tables_count'] for f in processed_files),
            'average_text_per_file': sum(f['text_length'] for f in processed_files) / len(processed_files) if processed_files else 0
        }
    }
    
    # Calculate file distribution by area committee and year
    distribution = {}
    for file_info in files_found:
        area = file_info['area_committee']
        year = file_info['year']
        
        if area not in distribution:
            distribution[area] = {}
        if year not in distribution[area]:
            distribution[area][year] = 0
        distribution[area][year] += 1
    
    summary['file_distribution'] = distribution
    
    # Save summary report
    summary_path = output_path / "conversion_summary.json"
    with open(summary_path, 'w', encoding='utf-8') as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    
    logger.info(f"Conversion completed. Summary saved to: {summary_path}")
    return summary

def main():
    parser = argparse.ArgumentParser(
        description='Bulk convert KrutiDev documents (PDF, DOCX, DOC) to Unicode with structured JSON output - macOS Version',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Bulk convert all files in a directory
  python bulk_krutidev_converter_macos.py --input-dir "/path/to/source" --output-dir "/path/to/output"
  
  # Convert with page-wise organization for PDFs
  python bulk_krutidev_converter_macos.py --input-dir "/path/to/source" --output-dir "/path/to/output" --pagewise
  
  # Convert without table extraction (faster)
  python bulk_krutidev_converter_macos.py --input-dir "/path/to/source" --output-dir "/path/to/output" --no-tables
  
  # Convert specific file types only
  python bulk_krutidev_converter_macos.py --input-dir "/path/to/source" --output-dir "/path/to/output" --extensions .pdf .docx

Required libraries for macOS:
  pip install PyPDF2 pdfplumber PyMuPDF pandas tabula-py python-docx
  brew install antiword pandoc openjdk
  pip install pypandoc docx2txt

Current User: vanshrajchauhan
Current Date: 2025-08-08 17:28:25 UTC
        """
    )
    
    parser.add_argument('--input-dir', help='Input directory containing nested files')
    parser.add_argument('--output-dir', help='Output directory for JSON files')
    parser.add_argument('--no-tables', action='store_true', help='Skip table extraction (faster processing)')
    parser.add_argument('--pagewise', action='store_true', help='Extract PDF content page-wise')
    parser.add_argument('--extensions', nargs='+', default=['.pdf', '.docx', '.doc'],
                       help='File extensions to process (default: .pdf .docx .doc)')
    
    args = parser.parse_args()
    
    if args.input_dir and args.output_dir:
        # Command line mode
        if not check_dependencies():
            return
        
        extract_tables = not args.no_tables
        
        print(f"Starting bulk conversion for macOS...")
        print(f"Current User: {CURRENT_USER}")
        print(f"Current Date: {CURRENT_DATE} {CURRENT_TIME} UTC")
        print(f"Input directory: {args.input_dir}")
        print(f"Output directory: {args.output_dir}")
        print(f"Extract tables: {extract_tables}")
        print(f"Page-wise extraction: {args.pagewise}")
        print(f"File extensions: {args.extensions}")
        
        summary = bulk_convert_files(
            args.input_dir, 
            args.output_dir, 
            extract_tables, 
            args.extensions,
            args.pagewise
        )
        
        print(f"\nConversion completed!")
        print(f"Files processed: {summary['conversion_summary']['successfully_processed']}")
        print(f"Files failed: {summary['conversion_summary']['failed_files']}")
        print(f"Total text extracted: {summary['statistics']['total_text_length']} characters")
        print(f"Total tables extracted: {summary['statistics']['total_tables']}")
        
    else:
        # Interactive mode
        print(f"Bulk KrutiDev to Unicode Converter - macOS Version")
        print(f"Current User: {CURRENT_USER}")
        print(f"Current Date: {CURRENT_DATE} {CURRENT_TIME} UTC")
        print("=" * 60)
        
        if not check_dependencies():
            return
        
        while True:
            print("\nOptions:")
            print("1. Bulk convert directory")
            print("2. Show directory structure preview")
            print("3. Test single file conversion")
            print("4. Check macOS dependencies")
            print("5. Quit")
            
            choice = input("\nSelect option (1-5): ").strip()
            
            if choice == '1':
                input_dir = input("Enter input directory path: ").strip()
                output_dir = input("Enter output directory path: ").strip()
                
                if not os.path.exists(input_dir):
                    print(f"Error: Input directory '{input_dir}' does not exist.")
                    continue
                
                extract_tables_choice = input("Extract tables? (y/n) [y]: ").strip().lower()
                extract_tables = extract_tables_choice != 'n'
                
                pagewise_choice = input("Use page-wise extraction for PDFs? (y/n) [n]: ").strip().lower()
                pagewise = pagewise_choice == 'y'
                
                extensions_input = input("File extensions to process (.pdf .docx .doc) [Enter for default]: ").strip()
                if extensions_input:
                    extensions = extensions_input.split()
                else:
                    extensions = ['.pdf', '.docx', '.doc']
                
                print(f"\nStarting bulk conversion...")
                print(f"This may take a while for large document collections...")
                
                summary = bulk_convert_files(input_dir, output_dir, extract_tables, extensions, pagewise)
                
                print(f"\nConversion completed!")
                print(f"Files processed: {summary['conversion_summary']['successfully_processed']}")
                print(f"Files failed: {summary['conversion_summary']['failed_files']}")
                print(f"Summary saved to: {Path(output_dir) / 'conversion_summary.json'}")
                
            elif choice == '2':
                input_dir = input("Enter directory path to preview: ").strip()
                if os.path.exists(input_dir):
                    files = find_files_recursively(input_dir, ['.pdf', '.docx', '.doc'])
                    print(f"\nFound {len(files)} files:")
                    
                    # Show distribution by area committee
                    distribution = {}
                    for file_info in files:
                        area = file_info['area_committee']
                        if area not in distribution:
                            distribution[area] = 0
                        distribution[area] += 1
                    
                    print("\nDistribution by Area Committee:")
                    for area, count in distribution.items():
                        print(f"  {area}: {count} files")
                    
                    # Show first few files as example
                    print(f"\nFirst 10 files:")
                    for i, file_info in enumerate(files[:10], 1):
                        print(f"  {i}. {file_info['relative_path']} "
                              f"(Area: {file_info['area_committee']}, Year: {file_info['year']})")
                    
                    if len(files) > 10:
                        print(f"  ... and {len(files) - 10} more files")
                else:
                    print(f"Directory '{input_dir}' does not exist.")
            
            elif choice == '3':
                file_path = input("Enter file path to test: ").strip()
                if os.path.exists(file_path):
                    file_info = {
                        'file_path': file_path,
                        'relative_path': Path(file_path).name,
                        'filename': Path(file_path).name,
                        'extension': Path(file_path).suffix.lower(),
                        'area_committee': 'test',
                        'year': 'test',
                        'path_parts': [],
                        'file_size': Path(file_path).stat().st_size
                    }
                    
                    result = process_single_file(file_info, True, False)
                    if result:
                        print("✅ File processed successfully!")
                        if 'content' in result:
                            print(f"Text length: {result['content']['text_length']}")
                            print(f"Tables found: {result['content']['tables_count']}")
                        else:
                            print(f"Text length: {result['summary']['total_text_length']}")
                            print(f"Tables found: {result['summary']['total_tables']}")
                    else:
                        print("❌ Failed to process file")
                else:
                    print(f"File '{file_path}' does not exist.")
            
            elif choice == '4':
                print("\nChecking macOS dependencies:")
                print(f"✅ Core libraries: Available")
                print(f"✅ PANDOC: {'Available' if PANDOC_AVAILABLE else 'Not available'}")
                print(f"✅ DOCX2TXT: {'Available' if DOCX2TXT_AVAILABLE else 'Not available'}")
                
                # Check antiword
                try:
                    result = subprocess.run(['antiword', '--version'], capture_output=True)
                    print(f"✅ Antiword: Available")
                except:
                    print(f"❌ Antiword: Not available (install with: brew install antiword)")
                
                # Check LibreOffice
                lo_path = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
                print(f"✅ LibreOffice: {'Available' if os.path.exists(lo_path) else 'Not available'}")
                
            elif choice == '5':
                print("Goodbye!")
                break
            else:
                print("Invalid option. Please try again.")

if __name__ == '__main__':
    main()